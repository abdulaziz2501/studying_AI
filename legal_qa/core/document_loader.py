"""
=============================================================
LEGAL QA SYSTEM — STEP 1 & 2: Document Loader & Preprocessor
=============================================================
Qo'llab-quvvatlanadigan formatlar: PDF, DOCX, TXT
Har bir hujjat metadata bilan birga yuklanadi
"""

import os
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional

# PDF o'qish uchun
import fitz  # PyMuPDF

# DOCX o'qish uchun
from docx import Document as DocxDocument


# -------------------------------------------------------
# Data Models
# -------------------------------------------------------

@dataclass
class DocumentChunk:
    """Hujjatning bir bo'lagi (chunk)"""
    text: str
    doc_name: str           # Fayl nomi
    doc_path: str           # To'liq yo'l
    page_number: int        # Sahifa raqami (PDF uchun)
    chunk_index: int        # Bu hujjatdagi necha-chi chunk
    section_hint: str = ""  # Bo'limga ishora (agar topilsa)
    char_start: int = 0     # Matnning boshlanish pozitsiyasi
    char_end: int = 0       # Matnning tugash pozitsiyasi


@dataclass
class LoadedDocument:
    """To'liq yuklangan hujjat"""
    name: str
    path: str
    file_type: str          # pdf / docx / txt
    full_text: str
    pages: List[str]        # Sahiyfa bo'yicha bo'lingan matn
    total_pages: int
    char_count: int
    word_count: int


# -------------------------------------------------------
# Text Cleaning
# -------------------------------------------------------

def clean_text(text: str) -> str:
    """
    Matndan keraksiz belgilarni tozalaydi.
    
    Args:
        text: Xom matn
    
    Returns:
        Tozalangan matn
    """
    # Ko'p bo'sh qatorlarni birlashtirish
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Ko'p bo'sh joylarni birlashtirish
    text = re.sub(r'[ \t]{2,}', ' ', text)
    
    # Har qator boshidagi/oxiridagi bo'shliqlarni olib tashlash
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Encoding xatolarini tuzatish (common replacements)
    text = text.replace('\x00', '')       # null bytes
    text = text.replace('\ufffd', ' ')    # replacement characters
    
    # Faqat bosib chiqarish mumkin bo'lgan belgilar + newline
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u0400-\u04FF\u00C0-\u024F]', ' ', text)
    
    return text.strip()


def extract_section_hint(text: str) -> str:
    """
    Matndan bo'lim nomini topishga urinadi.
    Masalan: 'Article 5', 'Section 3.2', 'CHAPTER I'
    
    Args:
        text: Chunk matni
    
    Returns:
        Bo'lim nomi yoki bo'sh string
    """
    patterns = [
        r'(?i)(article\s+\d+[\.\d]*)',
        r'(?i)(section\s+\d+[\.\d]*)',
        r'(?i)(chapter\s+[IVXivx\d]+)',
        r'(?i)(clause\s+\d+[\.\d]*)',
        r'(?i)(paragraph\s+\d+[\.\d]*)',
        r'(?i)(moddasi?\s+\d+[\.\d]*)',   # O'zbek: modda
        r'(?i)(bo\'lim\s+\d+[\.\d]*)',    # O'zbek: bo'lim
    ]
    
    first_200 = text[:200]
    for pattern in patterns:
        match = re.search(pattern, first_200)
        if match:
            return match.group(1).strip()
    
    return ""


# -------------------------------------------------------
# PDF Loader
# -------------------------------------------------------

def load_pdf(file_path: str) -> LoadedDocument:
    """
    PDF faylni sahifalab o'qiydi (PyMuPDF yordamida)
    
    Args:
        file_path: PDF fayl yo'li
    
    Returns:
        LoadedDocument obyekti
    """
    doc = fitz.open(file_path)
    pages_text = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text")
        cleaned = clean_text(page_text)
        pages_text.append(cleaned)
    
    doc.close()
    
    full_text = '\n\n'.join(pages_text)
    
    return LoadedDocument(
        name=Path(file_path).name,
        path=file_path,
        file_type="pdf",
        full_text=full_text,
        pages=pages_text,
        total_pages=len(pages_text),
        char_count=len(full_text),
        word_count=len(full_text.split()),
    )


# -------------------------------------------------------
# DOCX Loader
# -------------------------------------------------------

def load_docx(file_path: str) -> LoadedDocument:
    """
    DOCX faylni paragraflar bo'yicha o'qiydi
    
    Args:
        file_path: DOCX fayl yo'li
    
    Returns:
        LoadedDocument obyekti
    """
    doc = DocxDocument(file_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Jadval ichidagi matnlarni ham olish
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)
    
    # 50 ta paragrafdan 1 ta "sahifa" yasash
    pages_text = []
    chunk_size = 50
    for i in range(0, len(paragraphs), chunk_size):
        page_paragraphs = paragraphs[i:i + chunk_size]
        page_text = '\n'.join(page_paragraphs)
        pages_text.append(clean_text(page_text))
    
    if not pages_text:
        pages_text = [""]
    
    full_text = '\n\n'.join(pages_text)
    
    return LoadedDocument(
        name=Path(file_path).name,
        path=file_path,
        file_type="docx",
        full_text=full_text,
        pages=pages_text,
        total_pages=len(pages_text),
        char_count=len(full_text),
        word_count=len(full_text.split()),
    )


# -------------------------------------------------------
# TXT Loader
# -------------------------------------------------------

def load_txt(file_path: str) -> LoadedDocument:
    """
    TXT faylni o'qiydi (encoding aniqlash bilan)
    
    Args:
        file_path: TXT fayl yo'li
    
    Returns:
        LoadedDocument obyekti
    """
    # Har xil encodinglarni sinab ko'rish
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'cp1251']
    raw_text = None
    
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                raw_text = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    if raw_text is None:
        # So'nggi chora: xatolarni o'tkazib yuborish
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            raw_text = f.read()
    
    full_text = clean_text(raw_text)
    
    # Har 3000 belgidan bir sahifa
    pages_text = []
    page_size = 3000
    for i in range(0, len(full_text), page_size):
        pages_text.append(full_text[i:i + page_size])
    
    if not pages_text:
        pages_text = [""]
    
    return LoadedDocument(
        name=Path(file_path).name,
        path=file_path,
        file_type="txt",
        full_text=full_text,
        pages=pages_text,
        total_pages=len(pages_text),
        char_count=len(full_text),
        word_count=len(full_text.split()),
    )


# -------------------------------------------------------
# Main Folder Loader
# -------------------------------------------------------

def load_documents_from_folder(folder_path: str) -> List[LoadedDocument]:
    """
    Papkadagi barcha hujjatlarni yuklaydi.
    
    Args:
        folder_path: Hujjatlar papkasi yo'li
    
    Returns:
        LoadedDocument ro'yxati
    
    Raises:
        FileNotFoundError: Papka topilmasa
        ValueError: Hech qanday hujjat topilmasa
    """
    folder = Path(folder_path)
    
    if not folder.exists():
        raise FileNotFoundError(f"Papka topilmadi: {folder_path}")
    
    if not folder.is_dir():
        raise ValueError(f"Bu papka emas: {folder_path}")
    
    # Qo'llab-quvvatlanadigan formatlar
    supported_extensions = {
        '.pdf': load_pdf,
        '.docx': load_docx,
        '.txt': load_txt,
    }
    
    documents = []
    errors = []
    
    # Papkadagi barcha fayllarni ko'rib chiqish (pastki papkalar ham)
    all_files = []
    for ext in supported_extensions.keys():
        all_files.extend(folder.rglob(f"*{ext}"))
    
    all_files = sorted(all_files)  # Tartibli yuklash
    
    if not all_files:
        raise ValueError(
            f"Papkada hech qanday hujjat topilmadi: {folder_path}\n"
            f"Qo'llab-quvvatlanadigan formatlar: PDF, DOCX, TXT"
        )
    
    print(f"\n📂 Topilgan fayllar: {len(all_files)} ta")
    print("-" * 50)
    
    for file_path in all_files:
        ext = file_path.suffix.lower()
        loader_func = supported_extensions[ext]
        
        try:
            print(f"  ⏳ Yuklanmoqda: {file_path.name}...")
            doc = loader_func(str(file_path))
            documents.append(doc)
            print(f"  ✅ {file_path.name} — {doc.word_count} so'z, {doc.total_pages} sahifa")
        except Exception as e:
            error_msg = f"  ❌ {file_path.name}: {str(e)}"
            print(error_msg)
            errors.append(error_msg)
    
    if not documents:
        raise ValueError(
            f"Hech qanday hujjat muvaffaqiyatli yuklanmadi.\n"
            f"Xatolar: {chr(10).join(errors)}"
        )
    
    # Yakuniy hisobot
    total_words = sum(d.word_count for d in documents)
    print(f"\n📊 Yuklash yakunlandi:")
    print(f"   Muvaffaqiyatli: {len(documents)} ta")
    print(f"   Xato: {len(errors)} ta")
    print(f"   Jami so'zlar: {total_words:,}")
    
    return documents
